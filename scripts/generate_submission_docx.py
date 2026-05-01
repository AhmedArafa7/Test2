import base64
import os
import urllib.request
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

root = r"f:\GraduationProject"
out_dir = os.path.join(root, "deliverables")
img_dir = os.path.join(out_dir, "uml_images")
os.makedirs(img_dir, exist_ok=True)
out_doc = os.path.join(out_dir, "Baytology_Full_Requirements_and_UML_v3.docx")

def save_mermaid_png(source: str, out_path: str):
    encoded = base64.urlsafe_b64encode(source.encode("utf-8")).decode("utf-8")
    req = urllib.request.Request(
        f"https://mermaid.ink/img/{encoded}",
        headers={"User-Agent": "Mozilla/5.0"},
    )
    data = urllib.request.urlopen(req, timeout=60).read()
    with open(out_path, "wb") as f:
        f.write(data)

use_case = """
flowchart LR
Buyer[Buyer]-->A((Auth))
Buyer-->B((Browse Properties))
Buyer-->C((Save Property))
Buyer-->D((Create Booking))
Buyer-->E((AI Search))
Buyer-->F((AI Recommendations))
Buyer-->G((Chat))
Agent[Agent]-->A
Agent-->H((Manage Properties))
Agent-->G
Admin[Admin]-->I((Manage Users & Agents))
Admin-->J((Monitor Payments/Refunds))
Admin-->K((Monitor AI Requests))
PG[Payment Gateway]-->L((Webhook))
AIW[AI Worker]-->M((Resolve AI Callback))
"""

class_diagram = """
classDiagram
direction LR
class AppUser{
  +string Id
  +string Email
  +bool IsDeleted
}
class UserProfile{
  +Guid Id
  +string UserId
  +string DisplayName
  +string AvatarUrl
  +string PhoneNumber
  +Create()
  +Update()
}
class AgentDetail{
  +Guid Id
  +string UserId
  +string AgencyName
  +string LicenseNumber
  +decimal Rating
  +bool IsVerified
  +Create()
  +Update()
  +Verify()
}
class Property{
  +Guid Id
  +string AgentUserId
  +string Title
  +decimal Price
  +int Bedrooms
  +int Bathrooms
  +PropertyStatus Status
  +Create()
  +Update()
  +ChangeStatus()
  +AddImage()
}
class PropertyImage{
  +Guid Id
  +Guid PropertyId
  +string Url
  +bool IsPrimary
  +int SortOrder
}
class PropertyAmenity{
  +Guid Id
  +Guid PropertyId
  +bool HasParking
  +bool HasPool
  +bool HasGym
  +bool HasElevator
  +Update()
}
class SavedProperty{
  +Guid Id
  +string UserId
  +Guid PropertyId
  +DateTimeOffset SavedAt
}
class Booking{
  +Guid Id
  +Guid PropertyId
  +string UserId
  +string AgentUserId
  +Guid PaymentId
  +BookingStatus Status
  +Create()
  +Confirm()
  +Cancel()
}
class Payment{
  +Guid Id
  +Guid PropertyId
  +string PayerId
  +string PayeeId
  +decimal Amount
  +PaymentStatus Status
  +Create()
  +Complete()
  +MarkFailed()
}
class PaymentTransaction{
  +Guid Id
  +Guid PaymentId
  +string GatewayReference
  +string TransactionStatus
}
class Conversation{
  +Guid Id
  +Guid PropertyId
  +string BuyerUserId
  +string AgentUserId
  +DateTimeOffset LastMessageAt
  +Create()
  +SendMessage()
}
class Message{
  +Guid Id
  +Guid ConversationId
  +string SenderId
  +string Content
  +bool IsRead
  +Create()
  +MarkAsRead()
}
class SearchRequest{
  +Guid Id
  +string UserId
  +SearchInputType InputType
  +RequestStatus Status
  +Create()
  +Complete()
}
class SearchResult{
  +Guid Id
  +Guid SearchRequestId
  +Guid PropertyId
  +int Rank
}
class RecommendationRequest{
  +Guid Id
  +string RequestedByUserId
  +int TopN
  +RequestStatus Status
  +Create()
  +Complete()
}
class RecommendationResult{
  +Guid Id
  +Guid RequestId
  +Guid RecommendedPropertyId
  +int Rank
}
AppUser "1" --> "0..1" UserProfile : has
AppUser "1" --> "0..1" AgentDetail : has
AppUser "1" --> "0..*" Property : creates
Property "1" *-- "0..*" PropertyImage : composition
Property "1" *-- "0..1" PropertyAmenity : composition
AppUser "1" --> "0..*" SavedProperty : saves
Property "1" --> "0..*" SavedProperty : saved by
Property "1" --> "0..*" Booking : booked in
AppUser "1" --> "0..*" Booking : buyer
AppUser "1" --> "0..*" Booking : agent
Booking "1" --> "0..1" Payment : payment
Payment "1" *-- "0..*" PaymentTransaction : composition
Property "1" --> "0..*" Conversation : related
Conversation "1" *-- "0..*" Message : composition
SearchRequest "1" *-- "0..*" SearchResult : composition
RecommendationRequest "1" *-- "0..*" RecommendationResult : composition
"""

sequence = """
sequenceDiagram
actor Buyer
participant FE as Frontend
participant API as Baytology API
participant DB as Database
participant PG as Payment Gateway
participant NS as Notification Service
actor Agent
Buyer->>FE: Submit booking form
FE->>API: POST /bookings
API->>DB: Create Booking + Payment
API->>PG: Create payment intention
PG-->>FE: Checkout URL
PG-->>API: Webhook payment status
API->>DB: Update payment and booking status
API->>NS: Send notification
NS-->>Agent: Realtime notification
"""

use_case_png = os.path.join(img_dir, "use_case.png")
class_png = os.path.join(img_dir, "class_diagram.png")
sequence_png = os.path.join(img_dir, "sequence_diagram.png")
save_mermaid_png(use_case, use_case_png)
save_mermaid_png(class_diagram, class_png)
save_mermaid_png(sequence, sequence_png)

frs = [
    "FR-01: User Authentication and Account Management - The system shall allow users to register, log in, log out, confirm email, refresh access tokens, reset passwords, and delete accounts.",
    "FR-02: Role-Based Access Control - The system shall enforce authorization by role (Buyer, Agent, Admin) for both frontend routes and backend APIs.",
    "FR-03: Property Management (Agent) - The system shall allow Agents to create, update, and delete property listings and upload/manage property images.",
    "FR-04: Property Discovery - The system shall allow users to browse properties, view detailed property pages, and apply filters (city, district, listing type, price, area, etc.).",
    "FR-05: Saved Properties - The system shall allow Buyers to save and unsave properties for future reference.",
    "FR-06: Booking Management - The system shall allow Buyers to create bookings and allow Buyers/Agents to view and track booking details and status.",
    "FR-07: Payment Processing - The system shall support payment initiation for bookings and process payment gateway callbacks (webhooks) to update payment and booking states.",
    "FR-08: Refund Workflow - The system shall support refund request creation and admin-side review/processing.",
    "FR-09: Real-Time Chat - The system shall allow Buyers and Agents to create conversations and exchange real-time messages.",
    "FR-10: Notifications - The system shall generate and deliver real-time/in-app notifications (e.g., booking updates, payment updates, AI result completion).",
    "FR-11: AI Search - The system shall allow Buyers to submit AI-based search requests (text/voice/image + filters), monitor status, and retrieve ranked results.",
    "FR-12: AI Recommendations - The system shall allow Buyers to request AI-generated property recommendations and receive ranked recommendation results.",
    "FR-13: Admin Management Features - The system shall allow Admins to manage users/agents, monitor payments/refunds, oversee AI requests, and review audit logs.",
    "FR-14: Internal AI Integration - The system shall expose secure internal endpoints to receive asynchronous completion callbacks from AI worker services.",
]
nfrs = [
    "NFR-01: Security - The system shall implement secure authentication (JWT) and strict role-based authorization.",
    "NFR-02: Data Protection - The system shall use secure communication protocols (HTTPS, HSTS in production) and protect user/payment data.",
    "NFR-03: Performance - The system shall use caching and pagination to optimize read-heavy operations and reduce response times.",
    "NFR-04: Reliability - The system shall use asynchronous background processing and outbox/event mechanisms to ensure reliable message handling.",
    "NFR-05: Scalability - The system shall support service-level scalability through decoupled architecture and queue-based integrations.",
    "NFR-06: Availability and Responsiveness - The system shall provide real-time communication for chat and notifications with automatic reconnection support.",
    "NFR-07: Maintainability - The system shall follow a layered architecture (API, Application, Infrastructure, Domain) to improve maintainability and separation of concerns.",
    "NFR-08: Observability - The system shall provide structured logging and audit trails for troubleshooting, monitoring, and governance.",
    "NFR-09: Interoperability - The system shall expose RESTful APIs and development-time API documentation (OpenAPI/Swagger) for easier integration.",
    "NFR-10: Testability - The system shall be designed to support unit, integration, and workflow-level automated testing.",
]

doc = Document()
doc.add_heading("Baytology - Functional/Non-Functional Requirements and UML", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading("1. Functional Requirements", level=1)
for item in frs:
    doc.add_paragraph(item, style="List Bullet")
doc.add_heading("2. Non-Functional Requirements", level=1)
for item in nfrs:
    doc.add_paragraph(item, style="List Bullet")
doc.add_heading("3. Use Case Diagram", level=1)
doc.add_picture(use_case_png, width=Inches(6.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Figure 1 explains the main actor-to-system interactions in Baytology.").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading("4. Class Diagram", level=1)
doc.add_picture(class_png, width=Inches(6.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Figure 2 explains the core domain classes, members, and relationships.").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading("5. Sequence Diagram", level=1)
doc.add_picture(sequence_png, width=Inches(6.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Figure 3 explains the booking-payment workflow and realtime notification flow.").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save(out_doc)
print(out_doc)
